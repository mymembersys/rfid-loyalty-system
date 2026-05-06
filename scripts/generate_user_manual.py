"""Generate the RFID Loyalty System user manual as a .docx file.

Run: python scripts/generate_user_manual.py
Output: docs/RFID_Loyalty_System_User_Manual.docx
"""
from __future__ import annotations

import os
from datetime import date

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


# ---------- Theme ----------
PRIMARY = RGBColor(0x1F, 0x4E, 0x79)   # navy
ACCENT = RGBColor(0x2E, 0x75, 0xB6)    # blue
MUTED = RGBColor(0x6B, 0x72, 0x80)     # slate
PLACEHOLDER_BG = "EEF2F7"
PLACEHOLDER_BORDER = "C8D1DC"


# ---------- Helpers ----------
def shade_cell(cell, fill_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def set_cell_borders(cell, color_hex: str = PLACEHOLDER_BORDER, size: str = "8") -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color_hex)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = PRIMARY


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False,
             muted: bool = False, size: int = 11, align=None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if muted:
        run.font.color.rgb = MUTED


def add_bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(it)


def add_numbered(doc: Document, items: list[str]) -> None:
    for it in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(it)


def add_screenshot_placeholder(doc: Document, title: str, caption: str,
                               body_lines: list[str]) -> None:
    """A bordered, shaded box that stands in for a real screenshot."""
    table = doc.add_table(rows=2, cols=1)
    table.autofit = False
    table.columns[0].width = Cm(15.5)

    header = table.rows[0].cells[0]
    header.width = Cm(15.5)
    shade_cell(header, "1F4E79")
    set_cell_borders(header, "1F4E79", "12")
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run(f"  [Screen]  {title}")
    hr.bold = True
    hr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    hr.font.size = Pt(11)

    body = table.rows[1].cells[0]
    body.width = Cm(15.5)
    shade_cell(body, PLACEHOLDER_BG)
    set_cell_borders(body, PLACEHOLDER_BORDER, "8")
    body.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    body.paragraphs[0].text = ""
    for line in body_lines:
        p = body.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x33, 0x3B, 0x4A)

    doc.add_paragraph()
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(f"Figure: {caption}")
    cr.italic = True
    cr.font.size = Pt(9)
    cr.font.color.rgb = MUTED


def page_break(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    table.autofit = False
    table.columns[0].width = Cm(5)
    table.columns[1].width = Cm(10.5)
    for (k, v), row in zip(rows, table.rows):
        c0 = row.cells[0]
        c1 = row.cells[1]
        c0.width = Cm(5)
        c1.width = Cm(10.5)
        c0.text = ""
        c1.text = ""
        rk = c0.paragraphs[0].add_run(k)
        rk.bold = True
        rk.font.size = Pt(10)
        rv = c1.paragraphs[0].add_run(v)
        rv.font.size = Pt(10)


# ---------- Build ----------
def build() -> Document:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---------- Cover ----------
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("RFID Loyalty Card System")
    tr.bold = True
    tr.font.size = Pt(34)
    tr.font.color.rgb = PRIMARY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("User Manual")
    sr.font.size = Pt(20)
    sr.font.color.rgb = ACCENT

    doc.add_paragraph()
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = desc.add_run(
        "Multi-branch wellness loyalty platform — Admin Portal & Branch Terminal"
    )
    dr.italic = True
    dr.font.size = Pt(12)
    dr.font.color.rgb = MUTED

    for _ in range(8):
        doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(f"Version 1.0  ·  {date.today().isoformat()}")
    mr.font.size = Pt(11)
    mr.font.color.rgb = MUTED

    page_break(doc)

    # ---------- Table of contents ----------
    add_heading(doc, "Contents", level=1)
    toc = [
        "1. Introduction",
        "2. System overview",
        "3. Roles & access",
        "4. Getting started",
        "5. Admin Portal — Sign in",
        "6. Admin Portal — Dashboard",
        "7. Admin Portal — Members",
        "8. Admin Portal — Member Detail (cards, balances, voids)",
        "9. Admin Portal — Branches",
        "10. Admin Portal — Service Lines",
        "11. Admin Portal — Stamp Rules",
        "12. Admin Portal — Rewards Catalog",
        "13. Admin Portal — Reports & Export",
        "14. Admin Portal — Users",
        "15. Admin Portal — Settings (branding)",
        "16. Branch Terminal — Setup",
        "17. Branch Terminal — Check-in",
        "18. Branch Terminal — Redeem (stamps -> voucher)",
        "19. Branch Terminal — Voucher Redeem",
        "20. NFC & Balance QR",
        "21. Common workflows",
        "22. Troubleshooting",
        "23. Glossary",
    ]
    for t in toc:
        p = doc.add_paragraph(t)
        p.paragraph_format.left_indent = Cm(0.4)

    page_break(doc)

    # ---------- 1. Introduction ----------
    add_heading(doc, "1. Introduction", 1)
    add_para(doc,
             "The RFID Loyalty Card System is a card-only loyalty platform for a multi-branch "
             "wellness network covering diagnostic clinics, psychological clinics, and gyms. "
             "Members tap a physical RFID card at a branch terminal to record visits, accumulate "
             "stamps, and redeem rewards. There is no mobile app — everything happens at the "
             "front desk and in the admin portal.")
    add_para(doc, "This manual covers two applications:", bold=True)
    add_bullets(doc, [
        "Admin Portal (web) — used by HQ admins, branch managers, and auditors to manage "
        "members, cards, branches, rewards, stamp rules, and reports.",
        "Branch Terminal (web) — used by front-desk staff to record check-ins and process "
        "redemptions.",
    ])

    # ---------- 2. System overview ----------
    add_heading(doc, "2. System overview", 1)
    add_para(doc, "The platform consists of three components running together:")
    add_kv_table(doc, [
        ("API service",       "Node.js + Express REST API on port 4000."),
        ("Admin portal",      "React web app on port 5173 (admin@example.com)."),
        ("Branch terminal",   "React web app on port 5174 (frontdesk@example.com)."),
        ("Database",          "PostgreSQL 15+ (members, cards, visits, vouchers, audit log)."),
        ("RFID hardware",     "USB readers that emulate a keyboard — they 'type' the card UID and press Enter into the focused field."),
    ])

    # ---------- 3. Roles & access ----------
    add_heading(doc, "3. Roles & access", 1)
    add_para(doc, "Every user account is assigned exactly one role. The role controls which menu "
                  "items appear and which actions are allowed.")
    add_kv_table(doc, [
        ("admin",     "Full network access. Manages users, branches, service lines, settings, rewards, stamp rules."),
        ("manager",   "Branch-level management. Edits members, rewards and stamp rules; runs reports."),
        ("frontdesk", "Branch terminal only. Records visits and processes redemptions."),
        ("auditor",   "Read-only access to members, visits, and reports."),
    ])

    # ---------- 4. Getting started ----------
    add_heading(doc, "4. Getting started", 1)
    add_para(doc, "Default seeded credentials (development only — change before production):", bold=True)
    add_kv_table(doc, [
        ("HQ Admin",       "admin@example.com / admin123"),
        ("Branch Manager", "manager@example.com / manager123"),
        ("Front Desk",     "frontdesk@example.com / front123"),
    ])
    add_para(doc, "URLs:", bold=True)
    add_bullets(doc, [
        "Admin Portal: http://localhost:5173",
        "Branch Terminal: http://localhost:5174",
        "API base: http://localhost:4000",
    ])

    page_break(doc)

    # ---------- 5. Login ----------
    add_heading(doc, "5. Admin Portal — Sign in", 1)
    add_para(doc, "Open the admin portal URL in a browser. The sign-in card is displayed centred "
                  "on the page with the brand name and logo.")
    add_screenshot_placeholder(doc,
        "Sign-in screen",
        "Admin Portal sign-in page (/login).",
        [
            "+----------------------------------------------------+",
            "|              [LOGO]  RFID Loyalty                  |",
            "|         Sign in to manage members, cards,          |",
            "|                 and rewards.                       |",
            "|                                                    |",
            "|   Email     [ admin@example.com           ]        |",
            "|   Password  [ ********                    ]        |",
            "|                                                    |",
            "|             [        Sign in         ]             |",
            "+----------------------------------------------------+",
        ])
    add_para(doc, "Steps:", bold=True)
    add_numbered(doc, [
        "Enter your work email.",
        "Enter your password.",
        "Click Sign in. On success you are redirected to the Dashboard. On failure an inline red error appears under the password field.",
    ])

    # ---------- 6. Dashboard ----------
    add_heading(doc, "6. Admin Portal — Dashboard", 1)
    add_para(doc, "The Dashboard is the landing page after sign-in. It shows network-wide member "
                  "activity at a glance.")
    add_screenshot_placeholder(doc,
        "Dashboard",
        "Dashboard with five activity KPI cards.",
        [
            "Dashboard                                                ",
            "An overview of member activity across all branches.      ",
            "                                                         ",
            "+--------------+ +--------------+ +--------------+       ",
            "| Active (60d) | | Dormant(60d) | | Total active |       ",
            "|     412      | |      87      | |     499      |       ",
            "+--------------+ +--------------+ +--------------+       ",
            "+--------------+ +--------------+                        ",
            "| 2+ services  | | 3 services   |                        ",
            "|      63      | |      12      |                        ",
            "+--------------+ +--------------+                        ",
        ])
    add_para(doc, "KPI definitions:", bold=True)
    add_bullets(doc, [
        "Active (60d) — members with at least one visit in the last 60 days.",
        "Dormant (60d) — active members with no visit in the last 60 days.",
        "Total active — total number of members with active status.",
        "2+ services / 3 services — cross-service members who use multiple service lines.",
    ])

    page_break(doc)

    # ---------- 7. Members ----------
    add_heading(doc, "7. Admin Portal — Members", 1)
    add_para(doc, "The Members page lists every enrolled customer in the network, with quick "
                  "search and a New Member action.")
    add_screenshot_placeholder(doc,
        "Members list",
        "Searchable members table with the New Member button.",
        [
            "Members                                       [+ New Member]",
            "Enroll customers and manage their loyalty profiles.         ",
            "                                                            ",
            "[ Search name, member #, phone, email...     ] [ Search  ]  ",
            "                                                            ",
            "| Member #   | Name           | Email          | Status |   ",
            "|------------|----------------|----------------|--------|   ",
            "| M-000123   | Maria Cruz     | maria@...      | active |-> ",
            "| M-000124   | Juan Dela Cruz | juan@...       | active |-> ",
            "| M-000125   | Anna Reyes     | -              | susp.  |-> ",
        ])
    add_para(doc, "To enroll a new member:", bold=True)
    add_numbered(doc, [
        "Click + New Member (top-right).",
        "Fill in first/last name (required) plus optional email, phone, date of birth, gender, emergency contact.",
        "Pick the Origin branch — the branch that owns this member.",
        "Tick the marketing-consent checkbox if the member agreed.",
        "Click Create member. The list refreshes and the new member appears at the top.",
    ])
    add_screenshot_placeholder(doc,
        "Enroll New Member modal",
        "New Member dialog with required and optional fields.",
        [
            "+-----------------  Enroll New Member  ----------------+",
            "| First name * [ Maria       ]   Last name * [ Cruz  ]  |",
            "| Email        [ maria@...   ]   Phone        [ +63... ]|",
            "| DOB          [ 1992-03-04  ]   Gender       [ Female ]|",
            "| Emergency    [ Lina Cruz / +63...                    ]|",
            "| Origin branch[ Diagnostic — Manila                  v]|",
            "| [x] Member consents to marketing communications      |",
            "|                                                       |",
            "|                            [ Cancel ] [ Create member ]|",
            "+-------------------------------------------------------+",
        ])

    # ---------- 8. Member detail ----------
    add_heading(doc, "8. Admin Portal — Member Detail", 1)
    add_para(doc, "Click a member row to open the member detail page. This is the main workspace "
                  "for everything related to one member: profile, stamp balances, cards, visits, "
                  "and redemptions.")
    add_screenshot_placeholder(doc,
        "Member detail — header & balances",
        "Member detail page header with status, action buttons and per-service balances.",
        [
            "<- Members                                                  ",
            "                                                            ",
            "(MC)   Maria Cruz                                            ",
            "       M-000123 . [ active ] . joined 2025-08-10 . origin    ",
            "       Diagnostic-Manila                                     ",
            "                                                            ",
            "[ Suspend ] [ Blacklist ] [ Edit member ] [ Balance QR ]    ",
            "[ Write NFC URL ]                          [ + Issue Card ] ",
            "                                                            ",
            "+-- Profile -----------+  +-- Stamp balances -------------+ ",
            "| Email     maria@...  |  | Diagnostic   8                | ",
            "| Phone     +63...     |  | (5 earned . 0 spent)          | ",
            "| DOB       1992-03-04 |  | Gym          3                | ",
            "| ...                  |  |                               | ",
            "+----------------------+  +-------------------------------+ ",
        ])
    add_para(doc, "Cards section. Lists every card ever issued to this member with status "
                  "(active, replaced, lost, blocked). Active cards have Replace and Write NFC "
                  "actions on the right.")
    add_para(doc, "Visit history. Every check-in this member has made, with branch, service line, "
                  "stamps awarded, and a Void action for managers/admins.")
    add_para(doc, "Redemptions. Vouchers the member has earned with code, reward, status (pending, "
                  "redeemed, expired, voided), creation date, and expiry.")
    add_para(doc, "Common actions on this page:", bold=True)
    add_bullets(doc, [
        "+ Issue Card — bind a new RFID card UID to this member.",
        "Replace (per active card) — swap a lost/damaged card for a new UID; the old one is marked replaced.",
        "Edit member — change name, contact, status, marketing consent, origin branch.",
        "Suspend / Blacklist / Reactivate — change member status with audit trail.",
        "Balance QR — generate a 24-hour QR poster linking to the public balance page.",
        "Write NFC URL — burn a long-lived signed balance URL onto the card's NDEF record so a phone tap opens the balance page.",
        "Void (per visit) — mark a visit as voided with a required reason; visit stays in history but is excluded from reports.",
    ])

    page_break(doc)

    # ---------- 9. Branches ----------
    add_heading(doc, "9. Admin Portal — Branches", 1)
    add_para(doc, "Branches are the physical service locations in the network. Each branch belongs "
                  "to one service line.")
    add_screenshot_placeholder(doc,
        "Branches list",
        "Branches table with code, name, service line and status.",
        [
            "Branches                       [x] Show inactive  [+ New Branch]",
            "Service locations available in the network.                     ",
            "                                                                ",
            "| Code      | Name                  | Service Line | Status |   ",
            "| DG-MNL-01 | Diagnostic - Manila   | Diagnostic   | active |   ",
            "| PS-CEB-01 | Psych  - Cebu         | Psych        | active |   ",
            "| GY-DAV-01 | Gym - Davao           | Gym          | active |   ",
        ])
    add_para(doc, "To add a branch (admin only):", bold=True)
    add_numbered(doc, [
        "Click + New Branch.",
        "Enter Code (e.g. DG-MNL-02), Name, Service line, and optional address & phone.",
        "Click Create branch.",
    ])
    add_para(doc, "Deactivating a branch keeps existing visits, cards, and stamp rules intact, but "
                  "terminals signed in to that branch can no longer record new check-ins.",
                  italic=True, muted=True)

    # ---------- 10. Service lines ----------
    add_heading(doc, "10. Admin Portal — Service Lines", 1)
    add_para(doc, "Service lines define the high-level service categories (e.g. Diagnostic, Psych, "
                  "Gym). Only admins can edit them.")
    add_screenshot_placeholder(doc,
        "Service Lines",
        "Service Lines configuration with code, name, color and sort order.",
        [
            "Service Lines                  [x] Show inactive  [+ New Service Line]",
            "                                                                       ",
            "| Code        | Name        | Color    | Sort | Status |               ",
            "| diagnostic  | Diagnostic  | #1F4E79  |  10  | active |               ",
            "| psych       | Psychology  | #5B21B6  |  20  | active |               ",
            "| gym         | Gym         | #166534  |  30  | active |               ",
        ])
    add_para(doc, "Each service line carries a colour used to tint badges and tags across the "
                  "Admin Portal and Terminal. Codes must be lowercase a-z, digits and underscores.")

    # ---------- 11. Stamp rules ----------
    add_heading(doc, "11. Admin Portal — Stamp Rules", 1)
    add_para(doc, "Stamp rules tell the system how many stamps a visit awards and how often a "
                  "member can earn them. Branch-specific rules override the network default.")
    add_screenshot_placeholder(doc,
        "Stamp Rules list",
        "Stamp Rules with scope, cooldown and active window.",
        [
            "Stamp Rules                            [x] Show inactive [+ New Rule]",
            "Cooldown windows and stamps required per service line.               ",
            "                                                                     ",
            "| Service Line | Scope            | Stamps | Cooldown | Cross | ...  ",
            "| Diagnostic   | Network default  |   10   | 12 h     |  No   |      ",
            "| Gym          | Network default  |   8    | 12 h     |  Yes  |      ",
            "| Gym          | Gym - Davao      |   8    | 6 h      |  Yes  |      ",
        ])
    add_para(doc, "Field meaning:", bold=True)
    add_bullets(doc, [
        "Stamps required — number of stamps a member needs before they can claim a redemption.",
        "Cooldown (minutes) — minimum time between two paid stamps for the same member at the same service. 0 = no cooldown.",
        "Cross-service eligible — whether stamps from this rule can count toward any service line's reward.",
        "Active from / to — schedule rule activation; leave 'to' blank for no expiry.",
    ])

    page_break(doc)

    # ---------- 12. Rewards ----------
    add_heading(doc, "12. Admin Portal — Rewards Catalog", 1)
    add_para(doc, "The rewards catalog lists every item members can redeem with their stamps. "
                  "Each reward defines stamps cost, validity, and an optional per-member limit.")
    add_screenshot_placeholder(doc,
        "Rewards Catalog",
        "Rewards with code, stamp cost, validity and per-member limit.",
        [
            "Rewards Catalog                  [x] Show inactive  [+ New Reward]",
            "Items members can redeem with their stamps.                       ",
            "                                                                  ",
            "| Code     | Name                | Service | Stamps | Validity | ",
            "| GYM-DAY  | Free gym day pass   | Gym     |   8    | 30 days  | ",
            "| DG-15OFF | 15% off  diagnostic | Diag    |  10    | 60 days  | ",
            "| PS-FREE  | Free 30-min consult | Psych   |  12    | 90 days  | ",
        ])
    add_para(doc, "To create a reward:", bold=True)
    add_numbered(doc, [
        "Click + New Reward.",
        "Enter a unique Code (uppercase, e.g. GYM-DAY) and Name.",
        "Choose a Service line, or 'Any service' if it can be redeemed against any balance.",
        "Set Stamps cost, Validity (days the issued voucher is valid), and optionally Per-member limit.",
        "Click Create reward.",
    ])

    # ---------- 13. Reports ----------
    add_heading(doc, "13. Admin Portal — Reports & Export", 1)
    add_para(doc, "The Reports page combines a filter bar with five panels: a visits trend chart, "
                  "daily visits by branch, top members, pending vouchers, and redeemed vouchers.")
    add_screenshot_placeholder(doc,
        "Reports — filters & visits trend",
        "Filter bar plus the visits trend bar chart.",
        [
            "Reports                                  [ Reset ] [ Export to Excel ]",
            "Network activity at a glance.                                          ",
            "                                                                       ",
            "From [2026-04-05]  To [2026-05-05]  Branch [All v]  Reward [        ]  ",
            "                                                              [Apply] ",
            "                                                                       ",
            "Visits trend                                       412 total           ",
            "  ##   #####     ###    ##                                            ",
            "  ##   #####  ## ###  # ##  ###    Apr 5 ----------------- May 5      ",
        ])
    add_screenshot_placeholder(doc,
        "Reports — voucher tables",
        "Pending vouchers and redeemed vouchers with pagination.",
        [
            "Pending vouchers                                  3 outstanding         ",
            "| Voucher    | Reward     | Member        | Stamps | Created | Expires |",
            "| VC-AB12CD3 | GYM-DAY    | Maria Cruz    |   8    | Apr 30  | May 30  |",
            "| VC-EF45GH6 | DG-15OFF   | Juan Dela Cruz|  10    | May 1   | Jun 30  |",
            "                                                                        ",
            "  Rows per page [10 v]   Showing 1-3 of 3   << First  < Prev  1/1 >    ",
        ])
    add_para(doc, "Click Export to Excel to download an .xlsx with one sheet per panel: Summary, "
                  "Trend, Daily Visits, Top Members, Pending Vouchers, Redeemed Vouchers.")

    # ---------- 14. Users ----------
    add_heading(doc, "14. Admin Portal — Users", 1)
    add_para(doc, "Admins manage staff accounts under Users. Each user has email, full name, role, "
                  "an optional branch assignment, and an active flag.")
    add_screenshot_placeholder(doc,
        "Users",
        "Staff accounts with role and branch.",
        [
            "Users                                                [+ New User]",
            "                                                                  ",
            "| Email                  | Name        | Role     | Branch     | ",
            "| admin@example.com      | HQ Admin    | admin    | -          | ",
            "| manager@example.com    | Lina M.     | manager  | DG-MNL-01  | ",
            "| frontdesk@example.com  | Carlos R.   | frontdesk| GY-DAV-01  | ",
        ])
    add_para(doc, "Front-desk staff must be tied to a branch; admins and auditors are network-wide.")

    page_break(doc)

    # ---------- 15. Settings ----------
    add_heading(doc, "15. Admin Portal — Settings (branding)", 1)
    add_para(doc, "Settings let an admin customise the brand name, primary colour, accent colour, "
                  "and logo. The choices apply across the Admin Portal, Terminal, and the public "
                  "balance page.")
    add_screenshot_placeholder(doc,
        "Settings — branding",
        "Brand name, colour palette presets, hex inputs and logo upload.",
        [
            "Settings                                                            ",
            "                                                                    ",
            "Brand name [ RFID Loyalty                       ]                   ",
            "                                                                    ",
            "Quick palettes:                                                     ",
            " [ Office ] [ Indigo Pro ] [ Forest ] [ Ocean ]                     ",
            " [ Royal  ] [ Sunset    ] [ Rose   ] [ Slate  ]                    ",
            "                                                                    ",
            "Primary  [#1F4E79] []      Accent  [#2E75B6] []                     ",
            "                                                                    ",
            "Logo:    [ uploaded.png  (drag-and-drop or click to upload) ]      ",
            "                                                                    ",
            "                                              [ Save changes ]      ",
        ])
    add_para(doc, "Hex colours must be in #RRGGBB format. Logos are stored on the API server's "
                  "uploads folder.", muted=True)

    # ---------- 16. Terminal setup ----------
    add_heading(doc, "16. Branch Terminal — Setup", 1)
    add_para(doc, "The first time a tablet or PC opens the Terminal URL, it is unconfigured and "
                  "shows the Setup screen. Setup is a two-step process: sign in, then pick the "
                  "branch and service line.")
    add_screenshot_placeholder(doc,
        "Terminal — sign in",
        "Front-desk login at the start of terminal setup.",
        [
            "+----------------- [LOGO] RFID Loyalty -----------------+",
            "|  Terminal setup - sign in with a front-desk account   |",
            "|                                                       |",
            "|  Front-desk email  [ frontdesk@example.com         ]  |",
            "|  Password          [ ********                     ]  |",
            "|                                                       |",
            "|                          [        Sign in         ]  |",
            "+-------------------------------------------------------+",
        ])
    add_screenshot_placeholder(doc,
        "Terminal — branch & service",
        "After sign-in, pick the branch and service line for this terminal.",
        [
            "+--------- [LOGO] RFID Loyalty ---------+",
            "|  Signed in as Carlos R.               |",
            "|                                       |",
            "|  Branch       [ Gym - Davao        v]|",
            "|  Service line [ Gym                v]|",
            "|                                       |",
            "|              [ Start terminal -> ]    |",
            "+---------------------------------------+",
        ])
    add_para(doc, "The selection is saved in the browser's local storage so the terminal stays "
                  "configured across reboots. To reconfigure, click Sign out in the terminal "
                  "header.", muted=True)

    # ---------- 17. Check-in ----------
    add_heading(doc, "17. Branch Terminal — Check-in", 1)
    add_para(doc, "Check-in is the default mode after setup. The whole screen is dedicated to one "
                  "input field that auto-focuses, so a tap from a USB RFID reader fires submission "
                  "automatically.")
    add_screenshot_placeholder(doc,
        "Terminal — Check-in (idle)",
        "Idle Check-in screen waiting for a card tap.",
        [
            "+---- [LOGO] RFID Loyalty . Gym - Davao  [Gym] ----+    ",
            "|  [ Check-in ] [ Redeem ]  [ Voucher ]   Carlos R. |    ",
            "|--------------------------------------------------|    ",
            "|                                                  |    ",
            "|                Tap your card                     |    ",
            "|   Place the card on the reader to record         |    ",
            "|                  your visit.                     |    ",
            "|                                                  |    ",
            "|         [ Card UID                          ]    |    ",
            "|                          [ Record visit ]        |    ",
            "+--------------------------------------------------+    ",
        ])
    add_screenshot_placeholder(doc,
        "Terminal — Check-in (success)",
        "Welcome confirmation with stamp count and balance QR.",
        [
            "                Tap your card                            ",
            "                                                          ",
            "    Welcome, Maria Cruz!                                  ",
            "    Member #M-000123                                      ",
            "    Stamps: 9                                             ",
            "                                                          ",
            "                  +-----------+                           ",
            "                  |  [QR]     |   Scan with your phone    ",
            "                  +-----------+   to see your full balance",
        ])
    add_para(doc, "Notes:", bold=True)
    add_bullets(doc,
        [
            "USB RFID readers behave like a keyboard: they 'type' the UID and press Enter automatically. The clerk does not need to click anywhere.",
            "If the card is unknown or blocked, the green welcome panel is replaced by a large red error.",
            "Cooldown enforcement: if the same card has just earned a stamp for this service, the API rejects the second tap until the cooldown elapses. The visit still appears in the history but with 0 stamps.",
        ])

    page_break(doc)

    # ---------- 18. Redeem ----------
    add_heading(doc, "18. Branch Terminal — Redeem (stamps -> voucher)", 1)
    add_para(doc, "Redeem mode lets the front desk turn a member's stamps into a printable voucher "
                  "for a specific reward. The flow is: tap card -> see balances -> pick reward -> print.")
    add_screenshot_placeholder(doc,
        "Terminal — Redeem step 1 (tap)",
        "Tap a card to look up the member's balances.",
        [
            "+------- Redeem -------+                                    ",
            "| Tap the member's card to start.                          ",
            "|                                                          ",
            "| [ Card UID                              ]   [ Look up ]  ",
            "+----------------------------------------------------------+",
        ])
    add_screenshot_placeholder(doc,
        "Terminal — Redeem step 2 (select reward)",
        "Member balances and the reward catalogue side by side.",
        [
            "Maria Cruz . M-000123                              [Cancel]",
            "                                                            ",
            "Balances:   Diagnostic 8 . Psych 0 . Gym 9                  ",
            "                                                            ",
            "Available rewards (filtered to what is affordable):         ",
            "+-------------------+ +-------------------+                 ",
            "| GYM-DAY           | | DG-15OFF          |                 ",
            "| Free gym day pass | | 15% off diagnostic|                 ",
            "| Cost 8  Bal 9 OK  | | Cost 10  Bal 8 X  |                 ",
            "| [ Issue voucher ] | | [ Insufficient   ]|                 ",
            "+-------------------+ +-------------------+                 ",
        ])
    add_screenshot_placeholder(doc,
        "Terminal — Redeem step 3 (issued)",
        "Voucher issued with code, expiry and a print-friendly slip.",
        [
            "Voucher issued!                                              ",
            "                                                              ",
            "Code   VC-AB12CD3                                             ",
            "Reward Free gym day pass (GYM-DAY)                            ",
            "Stamps used 8 . Expires 2026-06-04                            ",
            "                                                              ",
            "        [ Print voucher ]   [ New redemption ]                ",
        ])
    add_para(doc, "Issued vouchers are also visible from the member detail page and the Reports "
                  "voucher tables.")

    # ---------- 19. Voucher Redeem ----------
    add_heading(doc, "19. Branch Terminal — Voucher Redeem", 1)
    add_para(doc, "Voucher Redeem is the second half of the redemption flow: the customer presents "
                  "an already-issued voucher (printed slip or shown on a phone) and the front desk "
                  "marks it consumed.")
    add_screenshot_placeholder(doc,
        "Terminal — Voucher lookup",
        "Look up a voucher by typing or scanning the voucher code.",
        [
            "+------- Voucher Redeem -------+                           ",
            "| Enter or scan the voucher code:                          ",
            "|                                                          ",
            "| [ VC-AB12CD3                        ]   [ Look up ]      ",
            "+----------------------------------------------------------+",
        ])
    add_screenshot_placeholder(doc,
        "Terminal — Voucher review & confirm",
        "Voucher details with confirm action when status is 'Ready to redeem'.",
        [
            "Voucher VC-AB12CD3                          [ Cancel ]      ",
            "                                                              ",
            "Reward Free gym day pass (GYM-DAY)                            ",
            "Member Maria Cruz . M-000123                                  ",
            "Stamps used 8                                                 ",
            "Issued 2026-05-04 . Expires 2026-06-04                        ",
            "                                                              ",
            "Status [ Ready to redeem ]                                    ",
            "                                                              ",
            "                          [   Confirm redemption   ]         ",
        ])
    add_para(doc, "Status indicators on the voucher view:", bold=True)
    add_bullets(doc, [
        "Ready to redeem — clerk can confirm; the voucher will move to redeemed.",
        "Already redeemed — voucher is consumed; no further action.",
        "Expired — voucher's expiry has passed; voiding it is a manager-level action elsewhere.",
        "Voided — voucher cancelled by an admin/manager; cannot be redeemed.",
    ])

    page_break(doc)

    # ---------- 20. NFC & QR ----------
    add_heading(doc, "20. NFC & Balance QR", 1)
    add_para(doc, "Members do not need a phone app to see their stamp balance. The system offers "
                  "two parallel mechanisms:")
    add_kv_table(doc, [
        ("Balance QR poster",
         "Generated from the member's profile page. It encodes a short-lived (24-hour) signed link to a public balance page. Use it on a printed card or as a one-off help to the customer."),
        ("NFC URL written to the card",
         "From the member detail page click Write NFC URL. The system generates a long-lived signed URL and the NfcWriter helper writes it to the card's NDEF URI record. Once written, the customer can tap the card on a phone to open the balance page directly."),
    ])
    add_para(doc, "After issuing a new card the system automatically prompts to write the NFC URL "
                  "so the card is fully provisioned in one motion.", muted=True)

    # ---------- 21. Common workflows ----------
    add_heading(doc, "21. Common workflows", 1)

    add_heading(doc, "21.1  Enrol a new member and issue a card", 2)
    add_numbered(doc, [
        "Admin Portal -> Members -> + New Member, fill the form, click Create member.",
        "Open the new member's profile.",
        "Click + Issue Card. Tap a blank card on the USB reader (or type the UID).",
        "Optional: pick the issuing branch (defaults to origin branch).",
        "Click Bind card. The Write NFC dialog opens automatically — tap the card on the NFC writer or click Skip.",
    ])

    add_heading(doc, "21.2  Replace a lost card", 2)
    add_numbered(doc, [
        "Open the member's profile.",
        "In the Cards table, click Replace on the lost card's row.",
        "Tap the new blank card. Click Replace card.",
        "The old card is marked 'replaced' (still in history); the new card becomes active.",
        "The Write NFC dialog opens — burn the URL onto the new card.",
    ])

    add_heading(doc, "21.3  Void a wrongly recorded check-in", 2)
    add_numbered(doc, [
        "Open the member's profile.",
        "In Visit history, click Void on the bad row.",
        "Type a reason (required; e.g. 'duplicate scan').",
        "Click Void visit. The visit stays in history with a 'voided' badge and is excluded from reports.",
    ])

    add_heading(doc, "21.4  Issue a reward voucher at the front desk", 2)
    add_numbered(doc, [
        "On the Terminal click Redeem.",
        "Tap the member's card.",
        "Pick the reward; only affordable rewards have an active Issue voucher button.",
        "Print the voucher (or note the code) and hand it to the customer.",
    ])

    add_heading(doc, "21.5  Customer comes back with a voucher", 2)
    add_numbered(doc, [
        "On the Terminal click Voucher.",
        "Type or scan the voucher code.",
        "Review reward, member and status. Status must be 'Ready to redeem'.",
        "Click Confirm redemption. Hand over the goods/service.",
    ])

    page_break(doc)

    # ---------- 22. Troubleshooting ----------
    add_heading(doc, "22. Troubleshooting", 1)
    add_kv_table(doc, [
        ("Card tap does nothing on terminal",
         "Click anywhere on the page — the UID input must be focused. The terminal auto-refocuses on every click."),
        ("'Card is replaced/blocked/lost'",
         "Card is no longer active. Issue a replacement from the member's profile in the Admin Portal."),
        ("'Member is suspended/blacklisted'",
         "Reactivate the member from the profile (admin/manager only)."),
        ("Cooldown rejection",
         "The same service line was stamped recently; check the rule's cooldown_minutes under Stamp Rules."),
        ("'Insufficient balance' on a redemption",
         "Member does not have enough stamps for the chosen reward. Cross-service rewards count any balance; service-specific rewards count only that service line's balance."),
        ("Reports panel shows empty",
         "Adjust the date range filters or click Reset. The default window is the last 30 days."),
        ("nfc_links table missing error",
         "Restart the API (it auto-creates the table on boot) or run npm run db:init from apps/api."),
        ("Terminal stuck after deactivating a branch",
         "Click Sign out on the terminal and run Setup again with a different branch."),
    ])

    # ---------- 23. Glossary ----------
    add_heading(doc, "23. Glossary", 1)
    add_kv_table(doc, [
        ("Branch",         "A physical service location (clinic or gym)."),
        ("Service line",   "A high-level service category — Diagnostic, Psych, Gym, etc."),
        ("Card",           "A physical RFID card bound to one member. UID is the card's unique identifier."),
        ("Visit",          "A recorded check-in. Awards stamps according to the matching stamp rule."),
        ("Stamp rule",     "Configuration controlling stamps awarded per visit and the cooldown window."),
        ("Stamp",          "Loyalty point earned per qualifying visit. Spent on rewards."),
        ("Reward",         "An item the member can claim with stamps. Has stamps_cost and validity_days."),
        ("Voucher",        "An issued claim against a reward. Codes look like VC-XXXXXXX."),
        ("Member status",  "active / suspended / blacklisted / inactive."),
        ("Cooldown",       "Minimum minutes between two stamp-earning visits for the same member at the same service."),
        ("Cross-service",  "When stamps from one service line can be spent on rewards belonging to another."),
        ("NDEF",           "The NFC data format used to write a URL onto a card so a phone tap opens it."),
    ])

    # ---------- Footer note ----------
    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run("End of manual.")
    fr.italic = True
    fr.font.color.rgb = MUTED
    fr.font.size = Pt(9)

    return doc


def main() -> None:
    here = os.path.dirname(os.path.abspath(__file__))
    out_dir = os.path.normpath(os.path.join(here, "..", "docs"))
    os.makedirs(out_dir, exist_ok=True)
    out = os.path.join(out_dir, "RFID_Loyalty_System_User_Manual.docx")

    doc = build()
    doc.save(out)
    print(f"Wrote: {out}")


if __name__ == "__main__":
    main()
